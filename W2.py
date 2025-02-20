import os
import docx
import time
import random  # For sample document generation (optional)

# --- Document Generation (Optional - for testing) ---
def create_sample_docx(filename, num_paragraphs=3):
    document = docx.Document()
    for _ in range(num_paragraphs):
        paragraph = document.add_paragraph()
        words = ["the", "quick", "brown", "fox", "jumps", "over", "lazy", "dog", "ipsum", "lorem", "search", "document", "text", "word", "python"]  # Added some keywords
        sentence = " ".join(random.choices(words, k=random.randint(5, 15))) + "."
        paragraph.add_run(sentence)
    document.save(filename)

# --- Text Extraction ---
def extract_text(docx_path, output_path):
    try:
        doc = docx.Document(docx_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
    except docx.opc.exceptions.PackageNotFoundError:
        print(f"Error: Could not open {docx_path}.  File not found or invalid format.")
    except Exception as e:
        print(f"Error extracting text from {docx_path}: {e}")

# --- Indexing ---
def build_index(extracted_dir):
    index = {}
    for filename in os.listdir(extracted_dir):
        if filename.endswith(".txt"):
            filepath = os.path.join(extracted_dir, filename)
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    text = f.read().lower()
                    for word in text.split():  # Simple word splitting (improve with regex later)
                        if word: # Skip empty strings
                            if word not in index:
                                index[word] = [filename]
                            else:
                                index[word].append(filename)
            except Exception as e:
                print(f"Error reading extracted text file {filepath}: {e}")

    return index

# --- Searching ---
def search(index, query):
    search_terms = query.lower().split()
    results = {}

    for term in search_terms:
        if term in index:
            for doc_id in index[term]:
                if doc_id not in results:
                    results[doc_id] = 0
                results[doc_id] += 1  # Simple ranking: count of matching terms

    # Sort results by relevance (number of matching terms)
    sorted_results = sorted(results.items(), key=lambda item: item[1], reverse=True)
    return sorted_results


# --- Main execution ---
if __name__ == "__main__":
    docx_dir = "docx_files"  # Directory for .docx files
    extracted_dir = "extracted_text"  # Directory for extracted .txt files

    # Create directories if they don't exist
    os.makedirs(docx_dir, exist_ok=True)
    os.makedirs(extracted_dir, exist_ok=True)

    # 1. Generate Sample Documents (Optional - Comment out if you have your own)
    for i in range(10):
        create_sample_docx(os.path.join(docx_dir, f"sample_{i}.docx"))


    # 2. Extract Text
    start_time = time.time()
    for filename in os.listdir(docx_dir):
        if filename.endswith(".docx"):
            docx_path = os.path.join(docx_dir, filename)
            txt_filename = filename[:-5] + ".txt"  # Change .docx to .txt
            txt_path = os.path.join(extracted_dir, txt_filename)
            extract_text(docx_path, txt_path)
    end_time = time.time()
    print(f"Extraction time: {end_time - start_time:.2f} seconds")


    # 3. Build Index
    start_time = time.time()
    index = build_index(extracted_dir)
    end_time = time.time()
    print(f"Index building time: {end_time - start_time:.2f} seconds")

    # 4. Search
    while True:
        query = input("Enter your search query (or type 'exit'): ")
        if query.lower() == "exit":
            break

        start_time = time.time()
        results = search(index, query)
        end_time = time.time()
        print(f"Search time: {end_time - start_time:.2f} seconds")


        if results:
            print("Results:")
            for doc_id, relevance in results:
                print(f"- {doc_id} (Relevance: {relevance})")
        else:
            print("No results found.")
